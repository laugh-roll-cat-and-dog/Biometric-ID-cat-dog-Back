#!/usr/bin/env python3
"""Generate CE68-GG-DEV-PROJECT_NAME.docx developer guide"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_document():
    doc = Document()

    # COVER PAGE
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CE68-GG-DEV-PROJECT_NAME\n\n')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Biometric Identification System for Dogs\n')
    run.font.size = Pt(16)
    run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Backend Developer Guide\n')
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Version: 1.0.0\nDate: April 2026\nFastAPI + PostgreSQL + PyTorch + YOLO')

    doc.add_page_break()

    # EXECUTIVE SUMMARY
    heading = doc.add_heading('Executive Summary', 0)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph('''This document provides comprehensive technical documentation for the Biometric Identification System backend. A FastAPI-based microservice designed to identify dogs using nose biometric embeddings.

System Components:
• FastAPI REST API with OpenAPI documentation
• PostgreSQL database with pgvector for embeddings
• PyTorch deep learning models (DINO/ConvNext)
• YOLO for nose detection
• Attention mechanisms (SAM, BAM, DAM)
• Image processing utilities

Key Features:
• Multi-image upload with automatic nose detection
• Dual embedding storage (full image + nose region)
• Real-time similarity search using cosine distance
• Attribute-based search across dog database
• Production-ready error handling
• CORS-enabled for cross-origin requests

Data Pipeline: Image Upload → Nose Detection → Embedding Generation → Database Storage → Similarity Search''')

    doc.add_page_break()

    # CHAPTER 1
    heading = doc.add_heading('Chapter 1: System Architecture', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    doc.add_heading('1.1 Components', level=2)
    doc.add_paragraph('''• Frontend/Client: HTTP request sender
• API Server: FastAPI application
• Database: PostgreSQL with pgvector
• ML Models: DINO, ConvNext, YOLO
• Storage: Filesystem for images''')

    doc.add_heading('1.2 Technology Stack', level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    cells = table.rows[0].cells
    cells[0].text = 'Component'
    cells[1].text = 'Technology'

    tech = [
        ('Framework', 'FastAPI 0.104+'),
        ('Database', 'PostgreSQL + pgvector'),
        ('ORM', 'SQLAlchemy 2.0+'),
        ('Deep Learning', 'PyTorch 2.0+'),
        ('Detection', 'YOLOv8'),
        ('Image Proc', 'OpenCV, Pillow')
    ]

    for i, (comp, tech_name) in enumerate(tech, 1):
        row = table.rows[i].cells
        row[0].text = comp
        row[1].text = tech_name

    doc.add_heading('1.3 Architecture Diagram', level=2)
    doc.add_paragraph('>>> INSERT YOUR MAIN ARCHITECTURE DIAGRAM HERE <<<')

    doc.add_page_break()

    # CHAPTER 2
    heading = doc.add_heading('Chapter 2: File & Module Structure', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    doc.add_heading('2.1 Root Level Files', level=2)
    doc.add_paragraph('app.py - FastAPI application factory', style='List Bullet')
    doc.add_paragraph('main.py - Alternative entry point', style='List Bullet')
    doc.add_paragraph('data.py - Dataset preparation utilities', style='List Bullet')
    doc.add_paragraph('cam_utils.py - GradCAM visualization', style='List Bullet')

    doc.add_heading('2.2 App Module Directory', level=2)
    table2 = doc.add_table(rows=8, cols=2)
    table2.style = 'Light Grid Accent 1'
    cells = table2.rows[0].cells
    cells[0].text = 'Module'
    cells[1].text = 'Contents'

    modules = [
        ('config/', 'database.py, settings.py'),
        ('models/', 'dog.py, dog_photo.py'),
        ('routes/', 'health.py, upload.py, search.py, searchbyimage.py'),
        ('utils/', 'embedding.py, crop.py, model_loader.py, file_handler.py, image_utils.py'),
        ('network/', 'network.py (Network_ConvNext)'),
        ('attention/', 'BAM.py, SAM.py, DAM.py, SEblock.py'),
        ('ai/', 'Model weights: .pt files'),
    ]

    for i, (module, content) in enumerate(modules, 1):
        row = table2.rows[i].cells
        row[0].text = module
        row[1].text = content

    doc.add_page_break()

    # CHAPTER 3
    heading = doc.add_heading('Chapter 3: Data Flow & Processing', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    doc.add_heading('3.1 Upload Pipeline', level=2)
    doc.add_paragraph('''Endpoint: POST /upload/photo
Input: images (list), name, breed, age, description, dog_id

Process:
1. Validate file format & size (max 8MB)
2. Create/retrieve Dog record
3. For each image:
   - Save to disk
   - Generate full-image embedding (1024-dim)
   - Detect nose using YOLO
   - Generate nose embedding
   - Store DogPhoto record

Response: 201 Created with file info

>>> INSERT UPLOAD FLOWCHART HERE <<<''')

    doc.add_heading('3.2 Search by Attributes', level=2)
    doc.add_paragraph('''Endpoint: POST /search
Input: query (value), search_mode (column)

Process:
1. Parse search_mode (name, breed, age, id, description)
2. Query Dog table
3. Fetch associated photos
4. Return top 3 results

>>> INSERT SEARCH FLOWCHART HERE <<<''')

    doc.add_heading('3.3 Search by Image', level=2)
    doc.add_paragraph('''Endpoint: POST /searchByImage
Input: image (query photo)

Process:
1. Upload image
2. Detect nose (YOLO)
3. Generate nose embedding
4. Calculate cosine similarity
5. Group by dog_id and average
6. Filter >= 0.7087 threshold
7. Return top 3 matches

>>> INSERT SIMILARITY SEQUENCE DIAGRAM HERE <<<''')

    doc.add_page_break()

    # CHAPTER 4
    heading = doc.add_heading('Chapter 4: Components & Functions', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    doc.add_heading('4.1 Database Models', level=2)
    doc.add_paragraph('''Dog Model:
• id (PK)
• name, breed, age
• description, created_at

DogPhoto Model:
• id (PK)
• dog_id (FK to Dog)
• filename, file_path, file_size_bytes
• file_extension, uploaded_at
• embedding (1024-dim vector)
• cropped_nose_path
• nose_embedding (1024-dim vector)''')

    doc.add_heading('4.2 API Endpoints', level=2)
    doc.add_paragraph('GET / - Test DB connectivity', style='List Bullet')
    doc.add_paragraph('GET /health - Health check', style='List Bullet')
    doc.add_paragraph('POST /upload/photo - Upload images', style='List Bullet')
    doc.add_paragraph('POST /search - Search by attributes', style='List Bullet')
    doc.add_paragraph('POST /searchByImage - Similarity search', style='List Bullet')

    doc.add_heading('4.3 Utility Functions', level=2)
    doc.add_paragraph('''embed_image(path: str) -> Tensor
  Generates 1024-dim embedding from image

crop_image(path: str, conf: float = 0.75) -> ndarray | None
  Detects and crops nose region

load_convnext_checkpoint_compat(model, path)
  Loads model weights with backward compatibility

detect_image_type(path: Path) -> str
  Detects image format

async save_file(file: UploadFile, dest: Path) -> int
  Streams file to disk (max 8MB)

to_public_image_path(path, dog_id, filename) -> str
  Converts filesystem path to public URL''')

    doc.add_heading('4.4 Neural Network', level=2)
    doc.add_paragraph('''Network_ConvNext(backbone: str, attention: str)
  
Backbone: DINO or V2 ConvNext Tiny
Attention: SAM, BAM, or both
Output: 1024-dim L2-normalized embedding

Flow: Image → Backbone → Attention → Pooling → Linear → LayerNorm → Embedding''')

    doc.add_heading('4.5 Attention Mechanisms', level=2)
    doc.add_paragraph('''SAM (Self-Attention): Spatial attention with learnable scaling''', style='List Bullet')
    doc.add_paragraph('''BAM (Bottleneck): Channel + spatial attention combined''', style='List Bullet')
    doc.add_paragraph('''DAM (Dual): Position + channel attention''', style='List Bullet')
    doc.add_paragraph('''SEblock: Squeeze-and-Excitation fusion''', style='List Bullet')

    doc.add_page_break()

    # CHAPTER 5
    heading = doc.add_heading('Chapter 5: API Specifications', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    doc.add_heading('5.1 Configuration', level=2)
    doc.add_paragraph('Base URL: http://localhost:8000', style='List Bullet')
    doc.add_paragraph('Swagger Docs: http://localhost:8000/docs', style='List Bullet')
    doc.add_paragraph('CORS: Allow all origins (*)', style='List Bullet')
    doc.add_paragraph('Max Upload: 8 MB', style='List Bullet')

    doc.add_heading('5.2 Response Examples', level=2)
    doc.add_paragraph('''Upload Success (201):
{
  "message": "Upload successful",
  "dog_id": 1,
  "total_images": 2,
  "uploaded_files": [...],
  "pet_info": {"name": "...", "breed": "...", ...}
}''')

    doc.add_heading('5.3 Error Handling', level=2)
    doc.add_paragraph('200 OK - Request successful', style='List Bullet')
    doc.add_paragraph('201 Created - Resource created', style='List Bullet')
    doc.add_paragraph('400 Bad Request - Invalid parameters', style='List Bullet')
    doc.add_paragraph('413 Payload Too Large - File > 8MB', style='List Bullet')
    doc.add_paragraph('500 Internal Server Error - Processing error', style='List Bullet')

    doc.add_page_break()

    # APPENDIX
    heading = doc.add_heading('Appendix', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    doc.add_heading('A. Key Dependencies', level=2)
    doc.add_paragraph('fastapi - Web framework', style='List Bullet')
    doc.add_paragraph('sqlalchemy - ORM', style='List Bullet')
    doc.add_paragraph('pgvector - Vector storage', style='List Bullet')
    doc.add_paragraph('torch - Deep learning', style='List Bullet')
    doc.add_paragraph('transformers - Models', style='List Bullet')
    doc.add_paragraph('opencv-python - Computer vision', style='List Bullet')
    doc.add_paragraph('pillow - Image processing', style='List Bullet')

    doc.add_heading('B. File Storage', level=2)
    doc.add_paragraph('''{BASE_UPLOAD_DIR}/
├── dog_1/
│   ├── uuid1.jpeg
│   └── nose/uuid1_nose.jpeg
├── dog_2/
│   ├── uuid2.jpeg
│   └── nose/uuid2_nose.jpeg''')

    doc.add_heading('C. Embeddings', level=2)
    doc.add_paragraph('Dimension: 1024-dim float32', style='List Bullet')
    doc.add_paragraph('Normalization: L2 normalization', style='List Bullet')
    doc.add_paragraph('Distance Metric: Cosine similarity', style='List Bullet')
    doc.add_paragraph('Threshold: 0.7087', style='List Bullet')
    doc.add_paragraph('Storage: PostgreSQL pgvector', style='List Bullet')

    doc.add_heading('D. Model Weights', level=2)
    doc.add_paragraph('dino_main_50_class.pt - Primary embedding model', style='List Bullet')
    doc.add_paragraph('yolo.pt - Nose detection model', style='List Bullet')
    doc.add_paragraph('v2_result.pt - ConvNext V2 weights', style='List Bullet')

    # Save file
    output_path = 'CE68-GG-DEV-PROJECT_NAME.docx'
    doc.save(output_path)
    
    file_size = os.path.getsize(output_path)
    print(f'✓ Document created successfully!')
    print(f'✓ Filename: {output_path}')
    print(f'✓ Size: {file_size / 1024:.1f} KB')
    print(f'✓ Location: {os.path.abspath(output_path)}')

if __name__ == '__main__':
    create_document()
