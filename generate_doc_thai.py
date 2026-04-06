#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_document():
    doc = Document()

    # หน้าปก (COVER PAGE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CE68-GG-DEV-PROJECT_NAME\n\n')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('ระบบระบุตัวตนทางชีวมิติสำหรับสุนัข\n(Biometric Identification System for Dogs)\n')
    run.font.size = Pt(16)
    run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('คู่มือสำหรับนักพัฒนาโปรแกรม\n')
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('เวอร์ชัน: 1.0.0\nวันที่จัดทำ: เมษายน 2026\nเทคโนโลยี: FastAPI + PostgreSQL + PyTorch + YOLO')

    doc.add_page_break()

    # บทสรุปผู้บริหาร (EXECUTIVE SUMMARY)
    heading = doc.add_heading('บทสรุปผู้บริหาร', 0)
    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph('''เอกสารฉบับนี้คือคู่มือสำหรับนักพัฒนาโปรแกรม (Developer Guide) ของระบบระบุตัวตนทางชีวมิติส่วนหลัง (Backend) ซึ่งพัฒนาด้วยสถาปัตยกรรม Microservice โดยใช้ FastAPI ทำหน้าที่ระบุตัวตนสุนัขอาศัยการประมวลผลการจดจำลายพิมพ์จมูก (Nose Biometric Embeddings)

ส่วนประกอบหลักของระบบมีดังนี้:
• FastAPI REST API สำหรับเซิร์ฟเวอร์
• ฐานข้อมูล PostgreSQL พร้อมไลบรารี pgvector สำหรับเก็บเวกเตอร์
• โมเดล Deep Learning ด้วย PyTorch (DINO/ConvNext)
• โมเดล YOLO สำหรับการตรวจจับตำแหน่งจมูก (Nose Detection)
• กลไกความสนใจ (Attention Mechanisms: SAM, BAM, DAM)
''')

    doc.add_page_break()

    # บทที่ 1
    heading = doc.add_heading('1. โครงสร้างของโปรแกรมภาพรวม (Top-Down Architecture)', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    doc.add_paragraph('ระบบถูกออกแบบในลักษณะ Top-Down โดยเริ่มจาก API Router รับ Request เข้ามา แล้วส่งไปประมวลผลที่ Utils โหลด Model จากนั้นบันทึกเก็บลง Database ความสัมพันธ์ระหว่างส่วนต่างๆ มีดังนี้:')
    
    doc.add_paragraph('>>> แทรกแผนภาพโครงสร้างสถาปัตยกรรมหลัก (ARCHITECTURE DIAGRAM) ที่นี่ <<<\n'
                      '(วางแผนภาพที่แสดงให้เห็นความสัมพันธ์ระหว่าง API Server, Database, รุ่น Model และ File Storage)')

    doc.add_heading('1.1 การแบ่งโมดูลหลักของระบบ', level=2)
    doc.add_paragraph('''ภายในโปรแกรมประกอบด้วย 1 โฟลเดอร์หลักคือ /app/ ซึ่งแบ่งออกเป็น 6 โมดูลย่อยที่ทำงานประสานกัน:
1. โมดูล Routes (/app/routes/): จัดการ API Endpoints คอยรับ Input (HTTP Requests) และส่ง Output (HTTP Responses)
2. โมดูล Models (/app/models/): จัดการโครงสร้างฐานข้อมูล (Database Schema) เชื่องโยง ORM ด้วย SQLAlchemy
3. โมดูล Utils (/app/utils/): ฟังก์ชันตัวช่วย (Utility Functions) จัดการไฟล์, ประมวลผลรูปภาพ, ครอปจมูกสุนัข
4. โมดูล Network (/app/network/): โครงสร้างสถาปัตยกรรม Neural Network สำหรับสร้าง Vector Embedding
5. โมดูล Attention (/app/attention/): รวมกลไกอัลกอริทึมของ Attention (BAM, SAM, DAM)
6. โมดูล Config (/app/config/): ตั้งค่าการเชื่อมต่อฐานข้อมูลและการตั้งค่าระบบพื้นฐาน''')

    doc.add_page_break()

    # บทที่ 2
    heading = doc.add_heading('2. รายละเอียดไฟล์และหน้าที่การทำงาน', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    doc.add_heading('2.1 ไฟล์ระดับ Root', level=2)
    doc.add_paragraph('โปรแกรมประกอบด้วยไฟล์ที่อยู่ระดับ Root ดังนี้:')
    doc.add_paragraph('1. app.py และ main.py: เป็นหน้าต่างหลักที่รวมแอปพลิเคชัน (FastAPI instance) มีหน้าที่ตั้งค่าระบบและเรียกใช้ Router ต่างๆ', style='List Bullet')
    doc.add_paragraph('2. data.py: จัดเตรียมข้อมูล (Dataset) สำหรับฝึกสอน YOLO Model', style='List Bullet')
    doc.add_paragraph('3. cam_utils.py: สร้างภาพฮีทแมพและ GradCAM เพื่อแสดงความสนใจของโมเดล', style='List Bullet')

    doc.add_heading('2.2 โมดูล /app/routes/ (การจัดการ API)', level=2)
    doc.add_paragraph('''ประกอบด้วยไฟล์จำนวน 4 ไฟล์ ทำหน้าที่เป็นโปรแกรมย่อยสำหรับจัดการ HTTP Request:
• health.py: ตรวจสอบสถานะการเชื่อมต่อ Database และ Server
• upload.py: หน้าที่รับไฟล์ภาพ โครอปรูป สร้าง Embedding และอัพโหลดเข้าฐานข้อมูล
• search.py: ค้นหาข้อมูลสุนัขจาก Attribute (ชื่อ, สายพันธุ์, รหัส)
• searchbyimage.py: ค้นหาสุนัขที่ตรงกับภาพอินพุตผ่านการเทียบค่า Cosine Similarity''')

    doc.add_heading('2.3 โมดูล /app/utils/ (ตัวช่วยและประมวลผล)', level=2)
    doc.add_paragraph('''ประกอบด้วยไฟล์จำนวน 5 ไฟล์:
• crop.py: ประมวลผลการตัดภาพจมูกสุนัขผ่านโมเดล YOLO
• embedding.py: สร้างฟีเจอร์เวกเตอร์ 1024-dim จากโมเดล DINO/ConvNext
• file_handler.py: ตรวจสอบและสตรีมไฟล์เพื่อบันทึกลงเครื่อง
• image_utils.py: แปลงผ่านของไฟล์ให้อยู่ในฟอร์แมต public URL
• model_loader.py: หน้าที่โหลดและดัดแปลงการเรียกใช้ Weight ใน Model (Backward compatibility)''')

    doc.add_heading('2.4 โมดูล /app/models/ (โครงสร้างฐานข้อมูล)', level=2)
    doc.add_paragraph('''• dog.py: ตารางเก็บข้อมูลทั่วไปของสุนัข (Dogs Table)
• dog_photo.py: ตารางเก็บข้อมูลรูปภาพที่ฝังเวกเตอร์และตำแหน่งความเชื่อมโยงกับสุนัข (DogPhotos Table)''')

    doc.add_page_break()

    # บทที่ 3
    heading = doc.add_heading('3. การทำงานของโปรแกรมย่อย (Input, Output และ Flow) ', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    doc.add_heading('3.1 โปรแกรมย่อยการอัพโหลดรูปภาพ (upload_photo)', level=2)
    doc.add_paragraph('''อธิบายการทำงาน: รับรูปภาพและข้อมูลพื้นฐานของสุนัขเข้ามา สร้างเวกเตอร์ทั้งภาพเต็มและภาพจมูกเพื่อบันทึกลง Database
• Input (พารามิเตอร์): 
  - images (List[UploadFile]): ไฟล์รูปภาพหลายไฟล์
  - name, breed, description (String): ข้อมูลทั่วไป
  - age, dog_id (Integer)
• Output (พารามิเตอร์ส่งกลับ): HTTP 201 Created ส่งกลับ ID สุนัข และจำนวนภาพที่อัพโหลดสำเร็จ''')

    doc.add_paragraph('\n>>> แทรก Flowchart หรือ แผนภาพขั้นตอนการทำงานกระบวนการอัพโหลดภาพ (UPLOAD FLOWCHART) ที่นี่ <<<')

    doc.add_heading('3.2 โปรแกรมย่อยการค้นหาด้วยแอตทริบิวต์ (search)', level=2)
    doc.add_paragraph('''อธิบายการทำงาน: ค้นหาข้อมูลสุนัขจากฟิลด์ต่างๆ ภายในฐานข้อมูล
• Input (พารามิเตอร์):
  - request (SearchRequest): คลาสประกอบด้วย `query` (คำที่ค้นหา) และ `search_mode` (ระบุคอลัมน์ เช่น breed, name, id)
• Output (พารามิเตอร์ส่งกลับ): HTTP 200 OK ข้อมูลสุนัขที่ตรงกัน 3 ผลลัพธ์และที่อยู่ภาพ''')

    doc.add_paragraph('\n>>> แทรก Flowchart หรือ แผนภาพขั้นตอนการสืบค้นข้อมูลปกติ (SEARCH FLOWCHART) ที่นี่ <<<')

    doc.add_heading('3.3 โปรแกรมย่อยการค้นหาด้วยรูปภาพ / คล้ายคลึง (search_by_image)', level=2)
    doc.add_paragraph('''อธิบายการทำงาน: รับภาพจมูกสุนัขที่ต้องการระบุตัวตน หาเวกเตอร์และหาค่าเฉลี่ยความคล้ายคลึงของเวกเตอร์ (Cosine distance) ทั้ง Database
• Input (พารามิเตอร์): 
  - image (UploadFile): ไฟล์รูปภาพหนึ่งไฟล์เพื่อคิวรีเทียบเคียง
• Output (พารามิเตอร์ส่งกลับ): HTTP 200 OK ส่งข้อมูลสุนัขที่คล้ายกันที่สุด 3 อันดับ (ค่า Threshold >= 0.7087) และความแม่นยำ (similarity score)''')

    doc.add_paragraph('\n>>> แทรก Sequence Diagram หรือ Flowchart การค้นหาด้วยใบหน้า/ภาพถ่าย (IMAGE SEARCH SEQUENCE/FLOWCHART) ที่นี่ <<<')

    doc.add_page_break()

    # บทที่ 4
    heading = doc.add_heading('4. รายละเอียดการทำงานของพารามิเตอร์และฟังก์ชันระดับลึก', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    doc.add_heading('4.1 ฟังก์ชัน crop_image()', level=2)
    doc.add_paragraph('''• พารามิเตอร์ Input: `image_path` (String/Path), `conf` (Float กำหนดเริ่มต้น = 0.75 สำหรับความเชื่อมั่นของ YOLO)
• การทำงาน: ส่ง `image_path` ไปที่โมเดล YOLO คัดลอกพิกัด Bbox ถ้าน้อยกว่า `conf` จะให้ค่า None หากสำเร็จจะตัดรูป (crop) ส่วนจมูกออก
• ส่งออก (Output): ตัวแปร `np.ndarray` ภาพส่วนที่มีเฉพาะจมูก หรือค่า `None` หากหาจำมูกไม่เจอ''')

    doc.add_heading('4.2 ฟังก์ชัน embed_image()', level=2)
    doc.add_paragraph('''• พารามิเตอร์ Input: `image` (String/Path หรือพาร์ทของไฟล์ภาพจมูก)
• การทำงาน: ทำภาพให้มีขนาด 224x224 (Resize) → แปลงเป็น Tensor → ส่งผ่านสถาปัตยกรรม Network_ConvNext (ที่มีโมดูล Attention เสริมอยู่)
• ส่งออก (Output): ลิสต์มิติเวกเตอร์ 1024 ค่าในรูปแบบตัวแปร `torch.Tensor` (L2 Normalized) เพื่อให้พร้อมสำหรับการจัดเก็บและหา Cosine Distance''')

    doc.add_page_break()

    # บทที่ 5
    heading = doc.add_heading('5. ข้อมูลจำเพาะและโมเดลเครือข่าย', 1)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    doc.add_paragraph('''5.1 โครงสร้างฐานข้อมูลตาราง DogPhoto
  • dog_id (Integer, Foreign Key ไปยังตาราง dogs)
  • filename (String 255)
  • file_path (String 500)
  • embedding (Vector ขนาด 1024 - เวกเตอร์ภาพเต็ม)
  • nose_embedding (Vector ขนาด 1024 - เวกเตอร์เฉพาะจมูก)''')

    doc.add_paragraph('''5.2 โครงสร้างโมเดลน้ำหนักเครือข่ายที่มีอยู่ (Model Checkpoints)
  • dino_main_50_class.pt : โมเดลสำหรับการสกัดคุณลักษณะ (Feature Extraction) หลัก
  • yolo.pt : โมเดลสำหรับดึงเฉพาะวัตถุเป้าหมาย (จมูกสุนัข)''')

    output_path = 'CE68-GG-DEV-PROJECT_NAME.docx'
    doc.save(output_path)
    
    print(f'✓ สร้างเอกสารภาษาไทยสำเร็จ!')
    print(f'✓ ชื่อไฟล์: {output_path}')
    print(f'✓ ขนาด: {os.path.getsize(output_path) / 1024:.1f} KB')

if __name__ == '__main__':
    create_document()
